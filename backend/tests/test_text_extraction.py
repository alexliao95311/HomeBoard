from pathlib import Path

from docx import Document as WordDocument
from openpyxl import Workbook

from app.services.text_extraction_service import (
    ExtractedSection,
    chunk_extracted_text,
    extract_document_text,
)


def test_extract_csv_preview(tmp_path: Path) -> None:
    path = tmp_path / "transactions.csv"
    path.write_text("date,vendor,amount\n2026-01-01,Water,125.50", encoding="utf-8")

    sections = extract_document_text(path, "text/csv")

    assert "date | vendor | amount" in sections[0].text
    assert "Water" in sections[0].text


def test_extract_docx_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "contract.docx"
    document = WordDocument()
    document.add_paragraph("Landscape service agreement")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Monthly fee"
    table.rows[0].cells[1].text = "$500"
    document.save(path)

    sections = extract_document_text(
        path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert "Landscape service agreement" in sections[0].text
    assert "Monthly fee | $500" in sections[0].text


def test_extract_xlsx_sheets_and_cells(tmp_path: Path) -> None:
    path = tmp_path / "budget.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Annual Budget"
    worksheet.append(["Category", "Amount"])
    worksheet.append(["Insurance", 12000])
    workbook.save(path)

    sections = extract_document_text(
        path,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    assert "[Sheet: Annual Budget]" in sections[0].text
    assert "Insurance | 12000" in sections[0].text


def test_chunk_text_uses_global_indexes_and_page_numbers() -> None:
    chunks = chunk_extracted_text(
        [
            ExtractedSection(text="A " * 1400, page_number=1),
            ExtractedSection(text="Second page", page_number=2),
        ]
    )

    assert [chunk.chunk_index for chunk in chunks] == list(range(len(chunks)))
    assert chunks[0].page_number == 1
    assert chunks[-1].page_number == 2
    assert all(len(chunk.text) <= 2500 for chunk in chunks)

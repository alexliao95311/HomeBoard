import csv
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path

import fitz
from docx import Document as WordDocument
from openpyxl import load_workbook

# Ligature characters that LaTeX/Type1 PDFs often leave un-expanded.
_LIGATURE_MAP = str.maketrans({
    "ﬀ": "ff",
    "ﬁ": "fi",
    "ﬂ": "fl",
    "ﬃ": "ffi",
    "ﬄ": "ffl",
    "ﬅ": "st",
    "ﬆ": "st",
    "’": "'",   # right single quotation → apostrophe
    "‘": "'",
    "“": '"',
    "”": '"',
    "–": "-",   # en dash
    "—": "--",  # em dash
    " ": " ",   # non-breaking space
})
_MULTI_BLANK_LINE_RE = re.compile(r"\n{3,}")


def _clean_text(text: str) -> str:
    """Expand ligatures, normalize Unicode, collapse excessive blank lines."""
    text = text.translate(_LIGATURE_MAP)
    text = unicodedata.normalize("NFKC", text)
    text = _MULTI_BLANK_LINE_RE.sub("\n\n", text)
    return text.strip()


@dataclass(frozen=True)
class ExtractedSection:
    text: str
    page_number: int | None = None


@dataclass(frozen=True)
class ExtractedChunk:
    text: str
    page_number: int | None
    chunk_index: int


class DocumentExtractionError(Exception):
    pass


def _extract_pdf(path: Path) -> list[ExtractedSection]:
    # TEXT_DEHYPHENATE removes soft hyphens at line ends.
    # Omitting TEXT_PRESERVE_LIGATURES causes PyMuPDF to expand fi/fl/ff etc.
    # sort=True reads blocks in natural reading order (helps with multi-column PDFs).
    _FLAGS = fitz.TEXT_DEHYPHENATE
    sections: list[ExtractedSection] = []
    with fitz.open(path) as pdf:
        for page_index, page in enumerate(pdf):
            raw = page.get_text("text", sort=True, flags=_FLAGS)
            sections.append(
                ExtractedSection(
                    text=_clean_text(raw),
                    page_number=page_index + 1,
                )
            )
    return sections


def _extract_docx(path: Path) -> list[ExtractedSection]:
    document = WordDocument(path)
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                lines.append(" | ".join(values))

    return [ExtractedSection(text=_clean_text("\n".join(lines)))]


def _extract_csv(path: Path) -> list[ExtractedSection]:
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as file:
        rows = [
            " | ".join(value.strip() for value in row)
            for row in csv.reader(file)
            if any(value.strip() for value in row)
        ]
    return [ExtractedSection(text="\n".join(rows))]


def _extract_xlsx(path: Path) -> list[ExtractedSection]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    sections: list[ExtractedSection] = []
    try:
        for worksheet in workbook.worksheets:
            rows = [f"[Sheet: {worksheet.title}]"]
            for row in worksheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    rows.append(" | ".join(values))
            sections.append(ExtractedSection(text="\n".join(rows)))
    finally:
        workbook.close()
    return sections


def extract_document_text(path: Path, content_type: str) -> list[ExtractedSection]:
    try:
        if content_type == "application/pdf":
            return _extract_pdf(path)
        if content_type == "text/csv":
            return _extract_csv(path)
        if (
            content_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            return _extract_docx(path)
        if (
            content_type
            == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ):
            return _extract_xlsx(path)
    except Exception as error:
        raise DocumentExtractionError("The document could not be read") from error

    raise DocumentExtractionError("Unsupported document type")


def _split_text(text: str, target_size: int = 2000, max_size: int = 2500) -> list[str]:
    remaining = text.strip()
    chunks: list[str] = []

    while remaining:
        if len(remaining) <= max_size:
            chunks.append(remaining)
            break

        split_at = remaining.rfind("\n", 1500, max_size + 1)
        if split_at < 1500:
            split_at = remaining.rfind(" ", 1500, max_size + 1)
        if split_at < 1500:
            split_at = min(target_size, len(remaining))

        chunk = remaining[:split_at].strip()
        if chunk:
            chunks.append(chunk)
        remaining = remaining[split_at:].strip()

    return chunks


def chunk_extracted_text(
    sections: list[ExtractedSection],
) -> list[ExtractedChunk]:
    chunks: list[ExtractedChunk] = []

    for section in sections:
        for text in _split_text(section.text):
            chunks.append(
                ExtractedChunk(
                    text=text,
                    page_number=section.page_number,
                    chunk_index=len(chunks),
                )
            )

    if not chunks:
        raise DocumentExtractionError("No extractable text was found")

    return chunks
